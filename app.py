import streamlit as st
import time
import base64

from amazon_web_scrape import *
from helpers_analytics import *
from helpers_ai_prompting import *
from user_login import *

from PIL import Image, UnidentifiedImageError
from docx import Document
from docx.shared import RGBColor  # Added for color control
from io import BytesIO
from datetime import datetime

from openai import OpenAI
from google import genai

from google.genai import types
from rainforest_api_credits import (
    get_credits,
    deduct_credits,
    MONTHLY_RAINFOREST_API_LIMIT,
)

import json
from pathlib import Path

import gspread
from google.oauth2.service_account import Credentials

RAINFOREST_API_KEY = st.secrets['RAINFOREST_API_KEY']
AMAZON_DOMAINS = ['amazon.com', 'amazon.co.uk']

# --- GOOGLE SHEETS LOGGING ---

_GS_SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]

@st.cache_resource
def get_log_sheet():
    """Return the first worksheet of the AmazonListingAI Google Sheet."""
    creds = Credentials.from_service_account_info(
        st.secrets["gcp_service_account"], scopes=_GS_SCOPES
    )
    client = gspread.authorize(creds)
    sheet = client.open(st.secrets["sheet"]["name"]).sheet1
    # Add header row if sheet is empty
    if sheet.row_count == 0 or sheet.cell(1, 1).value != "timestamp":
        sheet.insert_row(
            ["timestamp", "user", "subcategory", "function_name", "input_prompt", "output", "images_used"],
            index=1
        )
    return sheet

def log_to_sheets(function_name: str, input_prompt: str, output: str, images_used: int = 0):
    """Append a log row to the Google Sheet. Silently skips on error."""
    try:
        sheet = get_log_sheet()
        sheet.append_row([
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            st.session_state.get("current_user", "unknown"),
            st.session_state.get("product_selector", ""),
            function_name,
            input_prompt,
            output,
            images_used,
        ])
    except Exception as e:
        st.warning(f"⚠️ Logging failed: {e}")

# --- AUTHENTICATION ---

def check_authentication():
    """Check if user is authenticated, show login if not."""
    if 'authenticated' not in st.session_state:
        st.session_state['authenticated'] = False
    
    if not st.session_state['authenticated']:
        st.markdown("### 👋 欢迎光临")
        st.write("您好！请输入密码，开始管理您的亚马逊商品 ✨")
        
        password = st.text_input("Password", type="password", key="login_password")
        
        col1, col2, col3 = st.columns([1, 1, 2])
        with col1:
            if st.button("登录", type="primary"):
                if password in USER_PASSWORDS:
                    st.session_state['authenticated'] = True
                    st.session_state['current_user'] = USER_PASSWORDS[password]
                    st.rerun()
                else:
                    st.error("❌ Invalid password. Please try again.")
        
        return False  # Not authenticated - will block UI display
    else:
        # User is authenticated, show welcome message
        st.sidebar.success(f"✅ 欢迎 {st.session_state['current_user']}!")

        # Domain selector
        st.sidebar.selectbox(
            "Domain",
            options=AMAZON_DOMAINS,
            key="amazon_domain",
        )

        # Show monthly Rainforest API credits remaining
        try:
            credits_left = get_credits(st.session_state['current_user'])
            st.sidebar.metric(
                label="Rainforest API credits this month",
                value=f"{credits_left} / {MONTHLY_RAINFOREST_API_LIMIT}",
            )
        except Exception as e:
            st.sidebar.warning(f"⚠️ Could not load credits: {e}")

        if st.sidebar.button("Logout"):
            st.session_state['authenticated'] = False
            st.session_state.pop('current_user', None)
            st.rerun()

        return True  # Authenticated - allow UI display

def concatenate_input():
    full_listing = ""

    for i in range(1, 6):
        subheading = st.session_state.get(f"subheading{i}", "")
        bullet = st.session_state.get(f"bullet{i}", "")
        
        complete_bullet = subheading + ":" + bullet + "\n"
        full_listing += complete_bullet + "\n"
        
    return full_listing



def initialize_product_data(product_name, subcategory):
    """Initialize or update all session state variables for a product."""
    # Load raw data
    data, secondary_keywords = load_product_data(product_name, subcategory)
    
    # Update session state
    st.session_state['bullet_clusters'] = data['bullet_clusters']
    st.session_state['bullet_diagram'] = data['bullet_diagram'].fillna(0)
    st.session_state['bullet_labels'] = data['bullet_labels'].sort_values(by='monthly_sales', ascending=False)
    st.session_state['bullet_keywords'] = data['bullet_keywords']
    st.session_state['keyword_phrases'] = data['keyword_phrases']

    st.session_state['secondary_keywords'] = secondary_keywords
    
    # Update derived variables
    st.session_state['label_dict'] = get_cluster_dict(st.session_state['bullet_clusters'])
    st.session_state['labels'] = st.session_state['bullet_diagram'].index.tolist()
    st.session_state['best_phrases'] = {
        k: count_label_occurrences(
            st.session_state['label_dict'],
            st.session_state['bullet_labels'],
            k
        ) for k in st.session_state['labels']
    }


# --- Streamlit Setup ---
st.set_page_config(page_title="Image Keyword Filter", layout="wide")

# Get available products (loads in background even if not authenticated)
available_products = get_available_products()

if not available_products:
    st.error("No product directories found in data_files/")
    st.stop()

# Initialize default product if not set
if "product_selector" not in st.session_state:
    st.session_state["product_selector"] = available_products[0]

selected_product = st.session_state["product_selector"]

# Get subcategories for the selected product and initialize default
available_subcategories = get_product_subcategories(selected_product)

if not available_subcategories:
    st.error(f"No subcategory directories found in data_files/{selected_product}/")
    st.stop()

if "subcategory_selector" not in st.session_state or st.session_state.get("_last_product_for_subcategory") != selected_product:
    st.session_state["subcategory_selector"] = available_subcategories[0]
    st.session_state["_last_product_for_subcategory"] = selected_product

selected_subcategory = st.session_state["subcategory_selector"]

# Widget keys — must be set to "" so Streamlit re-renders them empty
_WIDGET_KEYS_TO_CLEAR = [
    "phrase_filter_and",
    "phrase_filter_or",
    "input_keywords",
    "listing_bullet_keywords",
    "keyword_listing_url_1",
    "keyword_listing_url_2",
    "keyword_listing_url_3",
    "product_specs",
    "finished_product_title",
]

# Non-widget session state keys — safe to delete outright
_STATE_KEYS_TO_DELETE = [
    "title_result",
    "ai_listing_draft_gpt",
    "ai_listing_draft_gemini",
    "photo_result",
    "suggestion_result",
    "synonym_result",
    "product_components_result",
    "rewriting_result",
    "listing_analysis",
    "formatted_search_terms",
    "product_listings_from_urls",
    "fixed_keywords",
    "filtered",
    "displayed_images",
    "previous_user_input",
    "previous_user_phrase",
    "previous_user_input_synonym",
    "product_description_result"
]

def _clear_product_state():
    """Reset all per-product session state so the UI shows a clean slate."""
    # Set widget-backed keys to "" so Streamlit renders them empty next pass
    for key in _WIDGET_KEYS_TO_CLEAR:
        st.session_state[key] = ""
    # Delete non-widget state outright
    for key in _STATE_KEYS_TO_DELETE:
        st.session_state.pop(key, None)

# Initialize or update data when product or subcategory changes (happens in background)
_current_selection = f"{selected_product}|{selected_subcategory}"
if "current_product" not in st.session_state or st.session_state["current_product"] != _current_selection:
    _clear_product_state()
    initialize_product_data(selected_product, selected_subcategory)
    st.session_state["current_product"] = _current_selection

    # Initialize displayed_images with first 24 rows
    initial_df = st.session_state['bullet_labels']
    st.session_state["displayed_images"] = initial_df.head(100) if len(initial_df) > 100 else initial_df.copy()

# --- Access session state variables (same names as before) ---
label_dict = st.session_state['label_dict']
labels = st.session_state['labels']
bullet_diagram = st.session_state['bullet_diagram']
best_phrases = st.session_state['best_phrases']
bullet_keywords = st.session_state['bullet_keywords']
keyword_phrases = st.session_state['keyword_phrases']
example_product_titles = "\n\n ".join(st.session_state['bullet_labels'].product_title.values.tolist()[:2])

# Apply colorization to diagram
bullet_diagram = colorize_df(bullet_diagram)

# Check authentication - only blocks UI display, not data loading
if not check_authentication():
    st.stop()  # Stop here if not authenticated - but data is already loaded!

# --- UI STARTS HERE (only shown if authenticated) ---
st.title("Amazon Listing Dashboard")

# Product + subcategory selectors side by side
_sel_col1, _sel_col2, _ = st.columns([2, 2, 6])
with _sel_col1:
    selected_product = st.selectbox(
        "Select Product Type",
        options=available_products,
        key="product_selector",
    )
with _sel_col2:
    selected_subcategory = st.selectbox(
        "Select Subcategory",
        options=available_subcategories,
        key="subcategory_selector",
    )

st.write()
 
left_col, _, right_col = st.columns([7.8, .8, 8])

with left_col:
    st.write("\n")
    st.dataframe(bullet_diagram, use_container_width=True)

with right_col:
    selected_label = st.selectbox("Select a category:", labels)

    st.markdown(f"#### {selected_label}")
    st.dataframe(best_phrases[selected_label], use_container_width=True, height=248)


def update_filter():
    label = st.session_state[f"label_input_{st.session_state['current_product']}"]  # CHANGE THIS LINE
    st.session_state["filtered"] = filter_bullets_by_phrase(
        st.session_state['bullet_labels'], label
    )
    # Update displayed images automatically on filter change
    DISPLAY_NUMBER = 100
    filter_result = st.session_state["filtered"]
    if len(filter_result) > DISPLAY_NUMBER:
        st.session_state["displayed_images"] = filter_result.head(DISPLAY_NUMBER)
    else:
        st.session_state["displayed_images"] = filter_result.copy()

with st.expander(f"竞品库 ({len(st.session_state['bullet_labels'])})"):

    st.text_input(
        "Enter Label",
        key=f"label_input_{st.session_state['current_product']}",  # CHANGE THIS LINE
        value=st.session_state.get("label_input", ""),  
        on_change=update_filter,
        width = 250
    )

    if "filtered" in st.session_state:
        filtered_df = st.session_state["filtered"]

    DISPLAY_NUMBER = 100

    # Render images from session state (outside button block so they persist)
    if "displayed_images" in st.session_state:
        trimmed_sample = st.session_state["displayed_images"]
        
        def to_str(val):
            """Convert list/ndarray to readable comma-separated string."""
            if isinstance(val, (list, set, tuple)):
                return ", ".join(map(str, val))
            if isinstance(val, np.ndarray):
                return ", ".join(map(str, val.tolist()))
            return str(val)

        # FIXED_SIZE = (1200, 1200)  # width, height for all images
        grid_cols = st.columns(3)

        for idx, (_, row) in enumerate(trimmed_sample.iterrows()):
            with grid_cols[idx % 3]:
                img_path = row["image_path"]
                try:
                    # Display image
                    st.image(img_path, use_container_width=True)
                    
                    # url
                    url = row['url']
                    st.caption(url)

                    # Extract bullet
                    full_listing = to_str(row.get("bullet_points", ""))

                    st.write(f"Monthly Sales: {to_str(row.get("monthly_sales", []))}")
                    st.markdown(
                        f"""
                        <div style="height: 88px; overflow: hidden;">
                            <strong>{to_str(row.get("product_title", []))}</strong>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                    st.write('')
                    with st.expander('Show full listing'):
                        st.markdown(row['bullet_points'].replace("\n", "\n\n"))

                    st.write("\n\n")

        

                

                except (FileNotFoundError, UnidentifiedImageError, OSError):
                    st.warning(f"⚠️ Could not load image: {img_path}")
           

# ---------------------------------------
#           OPENAI PHOTO FEATURE
# ---------------------------------------

@st.cache_resource
def get_openai_client(api_key):
    if "openai_client" not in st.session_state:
        st.session_state.openai_client = OpenAI(api_key=api_key)
    return st.session_state.openai_client

client = get_openai_client(st.secrets["OPENAI_KEY"])

@st.cache_resource
def get_gemini_client(api_key):
    if "gemini_client" not in st.session_state:
        st.session_state.gemini_client = genai.Client(api_key=api_key,
                                                      http_options=types.HttpOptions(timeout=100_000))
    return st.session_state.gemini_client

gemini_client = get_gemini_client(st.secrets["GEMINI_KEY"])

st.session_state.setdefault("photo_result", None)
st.session_state.setdefault("title_result", None)
st.session_state.setdefault("suggestion_result", None)
st.session_state.setdefault("synonym_result", None)
st.session_state.setdefault("product_components_result", None)
st.session_state.setdefault("rewriting_result", None)
st.session_state.setdefault("previous_user_input", None)
st.session_state.setdefault("previous_user_phrase", None)
st.session_state.setdefault("previous_user_input_synonym", None)
st.session_state.setdefault("ai_expander", False)
st.session_state.setdefault("default_tab", "Keywords & Title")
st.session_state.setdefault("pre_optimized_listing", False)
st.session_state.setdefault("finished_product_title", "")
st.session_state.setdefault("fixed_keywords", "")
st.session_state.setdefault("cleaned_keywords", "")
st.session_state.setdefault("formatted_search_terms", "")
st.session_state.setdefault("product_listings_from_urls", [])
st.session_state.setdefault("product_specs", "")
st.session_state.setdefault("product_reviews_summary", "")
st.session_state.setdefault("usage_keywords_from_reviews", "")
st.session_state.setdefault("combined_reviews_output", "")

st.session_state.setdefault("listing_analysis", "")
st.session_state.setdefault("product_info_synthesis", "")
st.session_state.setdefault("ai_listing_draft_gpt", "")
st.session_state.setdefault("ai_listing_draft_gemini", "")
st.session_state.setdefault("product_description_result", "")

st.session_state.setdefault("rainforest_asin_json", {})
st.session_state.setdefault("uploaded_images", [])


image_description_col, _, ai_tools_col = st.columns([5, 1, 8])

# =====================================================
# Image Description Generator
# =====================================================
@st.cache_data(show_spinner=False)
def load_images(files, max_size=1024):
    imgs = []
    for file in files:
        file.seek(0)
        img = Image.open(file).convert("RGB")
        img.thumbnail((max_size, max_size))
        imgs.append(img)
    return imgs

def sync_photo_results():
    """
    Called whenever uploaded_images_multiple changes.
    Clears out photo result so it doesn't show stale data.
    """
    st.session_state['photo_result'] = None

def process_multiple_images(images, client, system_prompt, user_instructions, model="gpt-4o-2024-08-06"):
    """
    Process multiple images in a single API call to OpenAI Vision.
    
    Args:
        images: List of PIL Image objects
        client: OpenAI client
        system_prompt: System message for the API
        user_instructions: Instructions for the user message
        model: OpenAI model to use (default: gpt-4o)
    
    Returns:
        str: Combined description from the API
    """
    # Convert images to base64
    image_messages = []
    for img in images:
        buffered = BytesIO()
        img.save(buffered, format="JPEG")
        img_base64 = base64.b64encode(buffered.getvalue()).decode()
        
        image_messages.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/jpeg;base64,{img_base64}"
            }
        })
    
    # Create messages with all images
    messages = [
        {
            "role": "system",
            "content": system_prompt
        },
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": user_instructions
                },
                *image_messages  # Unpack all image messages
            ]
        }
    ]
    
    # Make single API call with all images
    response = client.chat.completions.create(
        model=model, 
        messages=messages,
        max_completion_tokens=2000
    )
    
    return response.choices[0].message.content


with image_description_col:
    st.markdown("#### ✍️ Amazon Writing Assistant")

    with st.expander("Image Upload"):
        uploaded_images = st.file_uploader(
            "Upload multiple product photos",
            type=["jpg", "jpeg", "png"],
            key="uploaded_images_multiple",
            accept_multiple_files=True, 
            on_change=sync_photo_results
        )

        if not uploaded_images:
            st.session_state['uploaded_images'] = []
        else:
            st.session_state['uploaded_images'] = load_images(uploaded_images)
            # Create tab names: image_1, image_2, ...
            tab_names = [f"Image {i+1}" for i in range(len(st.session_state['uploaded_images']))]

            if len(tab_names) > 0:
                tabs = st.tabs(tab_names)

                # Put each image into its own tab
                for i, (tab, img) in enumerate(zip(tabs, st.session_state['uploaded_images'])):
                    with tab:
                        st.image(img, caption=f"Image {i+1}", use_container_width=True)


                if st.button("Generate description", key="generate_product_description_from_images"):
                    if len(st.session_state.get('uploaded_images') or []) > 3:
                        st.warning("请最多上传3张图片 (please have no more than 3 images)")
                    else:
                        # Progress bar
                        progress = st.progress(0)

                        # Process all images in a single API call
                        result = process_multiple_images(
                            st.session_state['uploaded_images'],
                            client,
                            "You are an expert at extracting product specifications and features from product images for Amazon listings",
                            image_specs_instructions
                        )

                        # Store result in session state
                        st.session_state['product_specs'] = result

                        # Complete progress
                        progress.progress(1.0)


    # Initialize product_specs in session_state if it doesn't exist
    if 'product_specs' not in st.session_state:
        st.session_state['product_specs'] = ""

    # Bind text_area to session_state
    facts_col, _, keyword_col = st.columns([7, 0.5, 6])
    with facts_col:
        # Specs
        st.text_area(
            "产品信息",
            key = 'product_specs',
            height=240,
            help="规格，特点，参数，尺寸，套数等等"
        )

        # Product URLs
        for i in range(1, 4):
            st.text_input(
                f"Example {i}",
                key=f"keyword_listing_url_{i}",
                placeholder="https://www.amazon.com/..."
            )

    with keyword_col:
        st.text_area("重点关键词", height=240, 
                     key='listing_bullet_keywords',
                     placeholder="e.g. gold picture frame, rustic ceramic vase",
                     help = "会写在listing每个卖点最前面")


    # Generate listing!
    st.write("")
    st.toggle(f"{st.session_state["amazon_domain"]}", key = 'rainforest_mode', value=True)
    st.write("")

    if st.button("Extract info"):
        
        # --- STEP 1: EXTRACT LISTING FROM URL ---
        product_urls = [
            st.session_state.get(f"keyword_listing_url_{i}", "").strip()
            for i in range(1, 4)
            if st.session_state.get(f"keyword_listing_url_{i}", "").strip()
        ]

        product_asins = list(set(
            asin for asin in 
            (extract_amazon_asin(url) for url in product_urls)
            if asin is not None
        ))

        # Check if product_listings is empty
        if len(st.session_state.get('uploaded_images') or []) > 3:
            st.warning("请最多上传3张图片 (please have no more than 3 images)")
        elif not st.session_state.get('uploaded_images'):
            st.warning("请先上传产品图")
        elif not st.session_state['product_specs']:
            st.warning("请先输入关产品规格")
        elif not st.session_state['listing_bullet_keywords']:
            st.warning("请先输入关键词 keywords")
        elif len(product_asins) == 0:
            st.warning("No valid asins provided")
        
        else:
            # Option a) Get locally
            if st.session_state['rainforest_mode'] == False:
                # reset reviews
                st.session_state['asin_reviews_str'] = ""

                bullet_labels = st.session_state['bullet_labels']
                st.session_state['product_listings_from_urls'] = bullet_labels[bullet_labels['ASIN'].isin(product_asins)]['bullet_points'].tolist()
            
            # Option b) Get via Rainforest API
            else:
                # Option b) Get via Rainforest API
                # ---- Credit check before making API calls ----
                current_user = st.session_state['current_user']
                credits_left = get_credits(current_user)

                if credits_left < len(product_asins):
                    st.warning(
                        f"You do not have enough credits remaining ({credits_left}). "
                        f"This request needs {len(product_asins)}."
                    )
                    
                else:
                    with st.spinner("Scraping URLs..."):
                        # ---- Rainforest API call ----
                        start_time=time.time()
                        st.session_state['rainforest_asin_json'] = get_amazon_product_data_parallel(
                            product_asins,
                            RAINFOREST_API_KEY,
                            domain=st.session_state["amazon_domain"],
                        )
                        end_time=time.time()
                        elapsed_time = end_time-start_time
                        st.write(f"Amazon request took {elapsed_time:.2f} seconds")

                        # Deduct credits for the requested ASINs
                        deduct_credits(current_user, len(product_asins))

                        if not st.session_state['rainforest_asin_json']:
                            st.warning("No response from Rainforest API call...consider using local mode")

                    # ---- Extract info from json payload ----
                    st.session_state['product_listings_from_urls'] = get_all_feature_bullets(st.session_state['rainforest_asin_json'])
                    st.session_state['asin_reviews_df'] = get_all_reviews(st.session_state['rainforest_asin_json'])
                    st.session_state['asin_reviews_str'] = concatenate_reviews(st.session_state['asin_reviews_df'])
                
                    combined_listing = "\n\n".join([f"Listing {i+1}:\n\n{bp}" for i, bp in enumerate(st.session_state['product_listings_from_urls'])])

                    log_to_sheets(
                        function_name="rainforest_api_listing",
                        input_prompt=", ".join(product_urls),
                        output=combined_listing,
                    )  

                    log_to_sheets(
                        function_name="rainforest_api_reviews",
                        input_prompt=", ".join(product_urls),
                        output=st.session_state["asin_reviews_str"],
                    ) 

                # ---- STEP 2: CHECK WHAT WE HAVE SO FAR (raw listings and reviews)
                # --- SCENARIO 1: NO REVIEWS OR FEATURE BULLETS:
                if not st.session_state['product_listings_from_urls'] and not st.session_state['asin_reviews_str']:
                    if st.session_state['rainforest_mode'] == True:
                        st.warning("请从竞品库选至少一个链接 url")
                    else:
                        st.warning("Sorry...please use local products for now.")

                else:
                    # --- SCENARIO 2: NO REVIEWS -------
                    if not st.session_state['asin_reviews_str']:
                        if not st.session_state.get('rainforest_mode', True):
                            st.warning("Could not extract reviews from JSON")
                        # --- STEP 1: ANALYZE FEATURE BULLETS ONLY -----
                        with st.spinner("Analyzing feature bullets..."):
                            combined_listing = "\n\n".join([f"Listing {i+1}:\n\n{bp}" for i, bp in enumerate(st.session_state['product_listings_from_urls'])])
                            listing_summary_prompt = listing_features_prompt_template.format(product_listings=combined_listing)
                            
                            st.session_state["product_info_synthesis"] = complete_phrase(
                                client,
                                listing_summary_prompt,
                                model='gpt-5.4-2026-03-05'
                            )

                            log_to_sheets(
                                function_name="analyze_listing",
                                input_prompt=", ".join(product_urls),
                                output=st.session_state["product_info_synthesis"],
                            )

                    # --- SCENARIO 3: HAS REVIEWS -------------
                    else:
                        # --- STEP 1: ANALYZE FEATURE BULLETS  -----
                        if not st.session_state['product_listings_from_urls']:
                            st.warning("Could not extract feature bullets from JSON...")
                            st.session_state["listing_analysis"] = ""
                        else:
                            with st.spinner("Analyzing feature bullets..."):
                                combined_listing = "\n\n".join([f"Listing {i+1}:\n\n{bp}" for i, bp in enumerate(st.session_state['product_listings_from_urls'])])
                                listing_summary_prompt = listing_features_prompt_template.format(product_listings=combined_listing)

                                st.session_state["listing_analysis"] = complete_phrase(
                                    client,
                                    listing_summary_prompt,
                                    model='gpt-5.4-2026-03-05'
                                )

                                log_to_sheets(
                                    function_name="analyze_listing",
                                    input_prompt=", ".join(product_urls),
                                    output=st.session_state["listing_analysis"],
                                )
                        
                        # --- STEP 2: ANALYZE REVIEWS  -----
                        with st.spinner("Analyzing reviews..."):
                            reviews_summary_prompt = reviews_summary_prompt_redacted_template.format(product_reviews=st.session_state['asin_reviews_str'])
                            reviews_usage_keywords_prompt = reviews_usage_keywords_prompt_template.format(product_reviews=st.session_state['asin_reviews_str'])
                            
                            st.session_state["product_reviews_summary"] = complete_phrase(
                                client,
                                reviews_summary_prompt,
                                model='gpt-5.4-2026-03-05'
                            )

                            st.session_state["usage_keywords_from_reviews"] = complete_phrase(
                                client,
                                reviews_usage_keywords_prompt,
                                model='gpt-5.4-2026-03-05'
                            )

                            st.session_state["combined_reviews_output"] = (
                                st.session_state["product_reviews_summary"]
                                + "\n\n"
                                + st.session_state["usage_keywords_from_reviews"]
                            )

                            log_to_sheets(
                                    function_name="analyze_reviews",
                                    input_prompt=", ".join(product_urls),
                                    output=st.session_state["combined_reviews_output"],
                                )
                        
                        # --- STEP 3: SYNTHESIZE INFO ----- 
                        with st.spinner("Synthesizing info..."):
                            listing_synthesizer_prompt = listing_synthesizer_prompt_template.format(listing_summary = st.session_state["listing_analysis"],
                                                                                            reviews_summary = st.session_state["combined_reviews_output"])

                            st.session_state["product_info_synthesis"] = complete_phrase(
                                client,
                                listing_synthesizer_prompt,
                                model='gpt-5.4-2026-03-05'
                            )
                
                    # --- STEP 3: WRITE LISTING FROM st.session_state["product_info_synthesis"]---
                    with st.spinner("Writing listing..."):
                        st.session_state['formatted_search_terms'] = "\n".join(
                            f"- {line}" for line in st.session_state["listing_bullet_keywords"].splitlines() if line.strip()
                        )

                        generate_listing_prompt_gpt = amazon_listing_prompt_template_revised_gpt.format(
                            product_specs=st.session_state['product_specs'],
                            keyword_search_phrases=st.session_state['formatted_search_terms'],
                            product_features=st.session_state["product_info_synthesis"]
                        )

                        generate_listing_prompt_gemini = amazon_listing_prompt_template_revised_gemini.format(
                            product_specs=st.session_state['product_specs'],
                            keyword_search_phrases=st.session_state['formatted_search_terms'],
                            product_features=st.session_state["product_info_synthesis"]
                        )

                        images = st.session_state.get('uploaded_images') or []

                        log_input = (
                            "specs:\n" + st.session_state.get('product_specs', '') +
                            "\nlisting analysis:\n" + st.session_state.get('product_info_synthesis', '') +
                            "\nkeywords:\n" + st.session_state.get('listing_bullet_keywords', '')
                        )

                        def run_gpt():
                            return complete_phrase(
                                client,
                                generate_listing_prompt_gpt,
                                model='gpt-5.4-2026-03-05',
                                images=images or None
                            )

                        def run_gemini():
                            gemini_listing_contents = [generate_listing_prompt_gemini] + images
                            return gemini_client.models.generate_content(
                                model="gemini-3-flash-preview",
                                contents=gemini_listing_contents,
                                config=types.GenerateContentConfig(
                                    thinking_config=types.ThinkingConfig(thinking_level="LOW"),
                                    temperature=1.0,
                                    http_options=types.HttpOptions(timeout=100_000),
                                )
                            ).text

                        start = time.time()

                        with ThreadPoolExecutor(max_workers=2) as executor:
                            future_gpt = executor.submit(run_gpt)
                            future_gemini = executor.submit(run_gemini)

                            # GPT result
                            try:
                                st.session_state["ai_listing_draft_gpt"] = future_gpt.result()
                                log_to_sheets(
                                    function_name="write_listing_draft_gpt",
                                    input_prompt=log_input,
                                    output=st.session_state["ai_listing_draft_gpt"],
                                    images_used=len(images),
                                )
                            except Exception as e:
                                st.session_state["ai_listing_draft_gpt"] = ""
                                st.warning(f"GPT listing generation failed: {e}")
                                log_to_sheets(
                                    function_name="write_listing_draft_gpt",
                                    input_prompt=log_input,
                                    output=f"ERROR: {e}",
                                    images_used=len(images),
                                )

                            # Gemini result
                            try:
                                st.session_state["ai_listing_draft_gemini"] = future_gemini.result()
                                log_to_sheets(
                                    function_name="write_listing_draft_gemini",
                                    input_prompt=log_input,
                                    output=st.session_state["ai_listing_draft_gemini"],
                                    images_used=len(images),
                                )
                            except Exception as e:
                                st.session_state["ai_listing_draft_gemini"] = ""
                                st.warning(f"Gemini listing generation failed: {e}")
                                log_to_sheets(
                                    function_name="write_listing_draft_gemini",
                                    input_prompt=log_input,
                                    output=f"ERROR: {e}",
                                    images_used=len(images),
                                )

                        elapsed = time.time() - start
                        st.write(f"Used {len(images)} images")
                        st.write(f"Request took {elapsed:.2f} seconds")


    if st.session_state["formatted_search_terms"]:
        st.write("#### Keywords")
        st.write(st.session_state["formatted_search_terms"])

    st.write("")
    tab1, tab2, tab3 = st.tabs([
        "Bullets Summary",
        "Reviews Summary",
        "Synthesized Info"
    ])

    with tab1:
        if st.session_state["listing_analysis"]:
            num_reference_listings = len(st.session_state['product_listings_from_urls'])
            st.write(f"##### Info from URLs ({num_reference_listings})")
            st.write(st.session_state["listing_analysis"])

    with tab2:
        if st.session_state["combined_reviews_output"]:
            st.write("##### Reviews summary")
            st.write(st.session_state["combined_reviews_output"])

    with tab3:
        if st.session_state["product_info_synthesis"]:
            st.write(f"##### Synthesized info ({len(st.session_state['product_listings_from_urls'])})")
            st.write(st.session_state["product_info_synthesis"])

    # Display results!
    def keyword_markdown(title, keyword_set):
        if not keyword_set:
            st.markdown(f"**{title}:** None ✅")
            return

        kw_string = ", ".join(sorted(keyword_set))
        st.markdown(f"**{title}:** {kw_string}")

    def display_listing_interface(listing, keyword_count_df, all_keywords, total_keywords_unique, new_words_added = None):
            st.write("")
            st.write(listing) 
            st.divider()

            stats_col1, _, stats_col2 = st.columns([7, 3, 5])
            with stats_col1:
                
                st.write(f"**Keywords provided**: {all_keywords}") 
                st.write(f"**Unique keywords in listing**: {total_keywords_unique}")
                if new_words_added:
                    keyword_markdown("New keywords added", new_words_added)
            
            with stats_col2:
                st.dataframe(keyword_count_df, hide_index=True, height = 176, width = 200)
       
# =====================================================
# Keywords / Writing Assistant
# =====================================================
with ai_tools_col:

    st.markdown("#### ")
    with st.expander("Title and keywords"):

        title_tab, condensor_tab = st.tabs(["Title", "Condense keywords"])
        
        with title_tab:

            col1, _, col2, _ = st.columns([4.2, 0.36, 5, 0.1])

            with col1:
                st.markdown("##### Title generator")

                # Text input for AND filtering (all words must be present)
                search_phrase_and = st.text_input(
                    "材质，尺寸，风格 etc.",
                    placeholder="E.g. ceramic, 5x7, 16",
                    key="phrase_filter_and"
                )
                
                # Text input for OR filtering (at least one word must be present)
                search_phrase_or = st.text_input(
                    "产品种类",
                    placeholder="E.g. tree, stem, mug cup",
                    key="phrase_filter_or"
                )
                
                # Apply filter based on inputs
                if (search_phrase_and and search_phrase_and.strip()) or (search_phrase_or and search_phrase_or.strip()):
                    # Filter the dataframe when there's input in either field
                    title_filtered_df = filter_by_phrase_final(
                        keyword_phrases,
                        input_phrase_and=search_phrase_and if search_phrase_and.strip() else None,
                        input_phrase_or=search_phrase_or if search_phrase_or.strip() else None, ignore_plural=False
                    )
                    title_filtered_df = title_filtered_df.sort_values(by = 'monthly_searches', ascending = False)
                    st.dataframe(title_filtered_df, height=352, hide_index=True)
                else:
                    st.dataframe(keyword_phrases, height=352, hide_index=True)
                
                # User enters keywords 
                st.text_area(
                    "Keywords",
                    placeholder="e.g. white ceramic vase, artificial hydrangea stems",
                    key="input_keywords",
                    height=240,
                    help = "AI 生成标题时会按搜索量排序"
                )

                user_input = st.session_state.get("input_keywords", "").strip()

                def generate_product_title():
                    user_input = st.session_state.get("input_keywords", "").strip()
                    if not user_input:
                        st.session_state["title_result"] = ""
                        return

                    # Block generation if more than 3 images uploaded
                    if len(st.session_state.get('uploaded_images') or []) > 3:
                        st.warning("请最多上传3张图片 (please have no more than 3 images)")
                        return

                    # If "Use uploaded images" is on but no images exist, block generation
                    if st.session_state.get('title_use_images', True) and not st.session_state.get('uploaded_images'):
                        st.warning("请先上传产品图")
                        return
                    
                    # Extract singular words + top phrases from user input
                    search_terms = user_input.split("\n")
                    sorted_search_terms = sort_search_terms(search_terms, keyword_phrases)
                    primary_keywords = extract_unique_words(sorted_search_terms)
                    top_search_terms = get_top_n_search_terms(sorted_search_terms)

                    # Writing the prompt
                    title_prompt = title_generator_prompt_gemini.format(
                                                        selected_product=selected_product,
                                                        top_search_terms=", ".join(top_search_terms),
                                                        primary_keywords=", ".join(primary_keywords),
                                                        secondary_keywords=st.session_state['secondary_keywords'],
                                                        example_product_titles=example_product_titles
                                                    )

                    start = time.time()
                    result = None

                    # --- Gemini attempts ---
                    try:
                        images = st.session_state.get('uploaded_images') or []
                        use_img = st.session_state.get('title_use_images', True)
                        gemini_contents = [title_prompt] + images if (use_img and images) else title_prompt

                        result = gemini_client.models.generate_content(
                            model="gemini-3-flash-preview",
                            contents=gemini_contents,
                            config=types.GenerateContentConfig(
                                thinking_config=types.ThinkingConfig(thinking_level="MINIMAL"),
                                max_output_tokens=100,
                                temperature=0.2,
                            )
                        ).text

                    except Exception as e:
                        st.warning(f"Using GPT backup...")

                    # --- GPT-5.1 fallback ---
                    if result is None:
                        # st.write("gpt backup")
                        title_prompt_gpt = title_generator_prompt_gpt.format(
                                                        selected_product=selected_product,
                                                        top_search_terms=", ".join(top_search_terms),
                                                        primary_keywords=", ".join(primary_keywords),
                                                        secondary_keywords=st.session_state['secondary_keywords'],
                                                        example_product_titles=example_product_titles
                                                    )
                        
                        for attempt in range(1, 4):
                            try:
                                if attempt > 1:
                                    wait = (attempt - 1) * 10
                                    st.warning(f"Retrying gpt (attempt {attempt})...")
                                    time.sleep(wait)

                                result = complete_phrase(
                                    client,
                                    title_prompt_gpt,
                                    model='gpt-5.4-2026-03-05',
                                    images=st.session_state.get('uploaded_images') or None
                                )
                                break  # Success — exit retry loop

                            except Exception as e:
                                if attempt == 3:
                                    st.error("Unfortunately, both Gemini and ChatGPT failed at this time. Please try again later.")

                    if result is not None:
                        st.session_state["title_result"] = result
                        log_to_sheets(
                            function_name="generate_title",
                            input_prompt=user_input,
                            output=result,
                            images_used=len(st.session_state.get('uploaded_images') or []),
                        )

                    end = time.time()
                    elapsed = end - start
                    st.write(f"Request took {elapsed:.2f} seconds")

                st.write("")
                use_images = st.toggle("参考上传图片", value=True, key="title_use_images")
                
                if st.button("Generate title"):
                    with st.spinner("Generating title..."):
                        generate_product_title()
            
            with col2:
                # Keyword prase
                st.write("#####")
                for i in range(8):
                    st.write('')
                st.markdown('场景词')
                st.dataframe(bullet_keywords[bullet_keywords.keyword.isin(st.session_state['secondary_keywords'])], height = 350, hide_index=True)

                # Display title
                if st.session_state["title_result"]:
                    st.write("")
                    st.write("")
                    st.markdown(f"**{st.session_state["title_result"]}**")
        
        with condensor_tab:
            st.text_area(
                    "Keywords",
                    placeholder="e.g. picture, frame, gold, ornate",
                    key="input_keywords_to_condense",
                    height=240, width= 360
                )
            
            if st.button("整理关键词"):
                st.session_state['cleaned_keywords'] = get_unique_words_from_string(st.session_state['input_keywords_to_condense'])
                st.write(st.session_state['cleaned_keywords'])


    if st.session_state["title_result"]:
        st.write("")
        st.write("")
        st.markdown(f"##### {st.session_state["title_result"]}")

    # Helper: generate a product description from a given listing draft.
    # Writes to the shared st.session_state["product_description_result"].
    def generate_description_for(listing_draft, source_label):
        primary_kw = ", ".join(st.session_state["formatted_search_terms"]) if isinstance(st.session_state["formatted_search_terms"], list) else st.session_state["formatted_search_terms"]
        secondary_kw = st.session_state.get("secondary_keywords", "")
        combined_keywords = f"{primary_kw}\n{secondary_kw}".strip()

        description_prompt = product_description_instructions.format(
            product_specs=st.session_state["product_specs"],
            keywords=combined_keywords,
            desirable_features=st.session_state["listing_analysis"],
            bullet_point_listing=listing_draft,
        )

        with st.spinner("Generating product description..."):
            st.session_state["product_description_result"] = complete_phrase(
                client,
                description_prompt,
                model='gpt-5.1-2025-11-13'
            )
            log_to_sheets(
                function_name=f"generate_product_description_{source_label}",
                input_prompt=(
                    "specs:\n" + st.session_state.get("product_specs", "") +
                    "\nkeywords:\n" + combined_keywords +
                    "\nlisting draft:\n" + listing_draft
                ),
                output=st.session_state["product_description_result"],
            )

    # Show GPT and Gemini drafts side-by-side in tabs.
    # Each tab gets its own "Generate product description" button, available
    # only when the respective draft exists.
    if st.session_state["ai_listing_draft_gpt"] or st.session_state["ai_listing_draft_gemini"]:
        gpt_tab, gemini_tab = st.tabs(["GPT", "Gemini"])

        with gpt_tab:
            if st.session_state["ai_listing_draft_gpt"]:
                st.write(st.session_state["ai_listing_draft_gpt"])
                st.write("")
                if st.button("Generate product description", key="generate_product_description_from_gpt"):
                    generate_description_for(
                        st.session_state["ai_listing_draft_gpt"],
                        source_label="gpt",
                    )

        with gemini_tab:
            if st.session_state["ai_listing_draft_gemini"]:
                st.write(st.session_state["ai_listing_draft_gemini"])
                st.write("")
                if st.button("Generate product description", key="generate_product_description_from_gemini"):
                    generate_description_for(
                        st.session_state["ai_listing_draft_gemini"],
                        source_label="gemini",
                    )

        # Download button lives outside the tabs and bundles whichever
        # drafts are available (GPT, Gemini, or both) plus the description.
        st.write("")

        def generate_docx():
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            doc = Document()

            def add_page_break():
                page_break_p = doc.add_paragraph()
                run = page_break_p.add_run()
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                run._element.append(br)

            def add_listing_content(content):
                """Render a listing draft (bullets / subheading:description lines) into the doc."""
                for line in content.split('\n'):
                    line = line.strip()
                    if not line:
                        continue

                    clean_line = line.replace("**", "")

                    if line.startswith(('-', '*')):
                        clean_line = clean_line.lstrip('-* ').strip()
                        p = doc.add_paragraph(style='List Bullet')
                    else:
                        p = doc.add_paragraph()

                    if ":" in clean_line:
                        subheading, description = clean_line.split(":", 1)
                        bold_run = p.add_run(subheading + ":")
                        bold_run.bold = True
                        p.add_run(description)
                    else:
                        p.add_run(clean_line)

            # 1. Title (shared across both drafts) if available
            title_text = st.session_state.get("title_result", "")
            if title_text:
                title_text = title_text.replace("**", "")
                heading = doc.add_heading(level=1)
                run = heading.add_run(title_text)
                run.font.color.rgb = RGBColor(0, 0, 0)
                doc.add_paragraph("")

            # 2. Drafts — include each available draft, separated by page breaks
            drafts = []
            if st.session_state.get("ai_listing_draft_gpt"):
                drafts.append(("GPT", st.session_state["ai_listing_draft_gpt"]))
            if st.session_state.get("ai_listing_draft_gemini"):
                drafts.append(("Gemini", st.session_state["ai_listing_draft_gemini"]))

            for i, (label, content) in enumerate(drafts):
                if i > 0:
                    add_page_break()
                section_heading = doc.add_heading(label, level=2)
                for run in section_heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                add_listing_content(content)

            # 3. Product description on a new page if available
            product_desc = st.session_state.get("product_description_result", "")
            if product_desc:
                add_page_break()
                desc_heading = doc.add_heading("Product Description", level=2)
                for run in desc_heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                for line in product_desc.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    clean_line = line.replace("**", "")
                    p = doc.add_paragraph()
                    if ":" in clean_line:
                        subheading, description = clean_line.split(":", 1)
                        bold_run = p.add_run(subheading + ":")
                        bold_run.bold = True
                        p.add_run(description)
                    else:
                        p.add_run(clean_line)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        st.download_button(
            label="Download Result",
            data=generate_docx(),
            file_name="Listing_Final.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    if st.session_state["product_description_result"]:
        st.write("#### Product Description")
        st.write(st.session_state["product_description_result"])