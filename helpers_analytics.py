import re
import streamlit as st
import numpy as np
import pandas as pd
import matplotlib as mpl

from pathlib import Path
from collections import Counter

from datetime import datetime
from docx import Document
from docx.shared import Pt
import io

from typing import Optional
from urllib.parse import urlparse, parse_qs

STOP_WORDS = {
    "a","an","the","and","or","but","if","while","with","without","of","for","to","from",
    "in","on","at","by","up","down","off","over","under","into","onto","as","is","are",
    "was","were","be","been","being","that","this","these","those","it","its","it's",
    "your","my","our","their","his","her","they","them","he","she","we","you","i",
    "so","not","no","yes","can","could","should","would","will","just","than","then",
    "about","all","any","some","more","most","other","such","only","own","same","too",
    "very","both","each","few","many","much","every","also","here","there","again",
    "once","because","why","how","what","when","where","which","who","whom",
    "do","does","did","doing","done","has","have","had","having",
    "am","shall","may","might","must",
}

def extract_paragraphs_by_phrase(text: str, phrase: str):
    """
    Return a list of all paragraphs/lines that contain `phrase`.
    A "paragraph" is defined as the text from a line start (or previous newline)
    up to the next newline. Matching is case-insensitive.
    """
    if not phrase:
        return []

    # escape phrase (so user input like "vintage (design)" won't break the regex)
    esc = re.escape(phrase)
    # pattern: start at line boundary (^ or \n), capture any chars except newline that include the phrase
    pattern = rf'(?im)(?:^|\n)([^\n]*\b{esc}\b[^\n]*)'
    matches = re.findall(pattern, text)
    # strip extra whitespace and return
    return [m.strip() for m in matches]


def extract_first_paragraph_by_phrase(text: str, phrase: str):
    """
    Return the first paragraph/line that contains `phrase` or None if not found.
    """
    res = extract_paragraphs_by_phrase(text, phrase)
    return res[0] if res else None


def colorize_df(df: pd.DataFrame, base_cmap: str = "Greens", min_frac: float = 0.05, max_frac: float = 1.5):
    # Get original colormap
    cmap_orig = mpl.colormaps[base_cmap]
    
    # Slice colormap to soften extremes
    cmap_vals = cmap_orig(np.linspace(min_frac, max_frac, 256))
    # Remove alpha channel for pandas
    cmap_vals = cmap_vals[:, :-1]
    soft_cmap = mpl.colors.ListedColormap(cmap_vals)
    
    return (
        df.style
        .background_gradient(cmap=soft_cmap, vmin=0, vmax=1.0, axis=None)  # Changed to 0 and 1.0
        .format("{:.2f}")
    )

def get_cluster_dict(picture_cluster_df_labeled):
    """
    Create a dictionary mapping each cluster label to a list of unique keywords sorted alphabetically.
    """
    cluster_dict = picture_cluster_df_labeled.groupby("label")["keywords"].apply(list).to_dict()
    
    # Flatten all keyword lists per label, remove duplicates, and sort
    cluster_dict = {
        k: sorted(set([item for sublist in v for item in sublist]))
        for k, v in cluster_dict.items()
    }
    
    return cluster_dict


def filter_bullets_by_phrase(df, phrase):
    """
    Filter dataframe rows by counting keyword matches in product_title and parsed_headings.
    
    Args:
        df: DataFrame with 'product_title' (str) and/or 'parsed_headings' (list of str) columns
        phrase: Search phrase string (e.g., 'ornate gold frame')
    
    Returns:
        DataFrame with rows where match count > 0, sorted by match count (descending)
    """
    if df.empty or not phrase or not phrase.strip():
        return df.iloc[0:0]  # Return empty DataFrame with same structure
    
    # Check which columns exist
    has_title = 'product_title' in df.columns
    
    if not has_title:
        return df.iloc[0:0]  # Return empty DataFrame if neither column exists
    
    # Tokenize search phrase (lowercase for case-insensitive matching)
    keywords = phrase.lower().split()
    
    # Vectorized counting function
    def count_matches(row):
        count = 0
        
        # Count matches in product_title
        if has_title:
            title_lower = str(row['product_title']).lower()
            count += sum(keyword in title_lower for keyword in keywords)
        
            
        return count
    
    # Apply counting and filter
    match_counts = df.apply(count_matches, axis=1)
    filtered_df = df[match_counts > 0].copy()
    
    # Add match count column and sort
    if not filtered_df.empty:
        filtered_df['_match_count'] = match_counts[match_counts > 0]
        filtered_df = filtered_df.sort_values(['_match_count', 'monthly_sales'], ascending=False)
    
    return filtered_df





# --- Helper Functions ---
def get_available_products(data_dir="data_files"):
    """Get list of product directories in data_files."""
    data_path = Path(data_dir)
    if not data_path.exists():
        return []
    return sorted([d.name for d in data_path.iterdir() if d.is_dir() and not d.name.endswith('checkpoints')])

def get_product_subcategories(product_name, data_dir="data_files"):
    """Get list of subcategory directories for a given product."""
    product_path = Path(data_dir) / product_name
    if not product_path.exists():
        return []
    return sorted([d.name for d in product_path.iterdir() if d.is_dir() and not d.name.endswith('checkpoints')])

@st.cache_data
def load_product_data(product_name, subcategory, data_dir="data_files"):
    """Load all parquet files for a specific product subcategory by finding files with matching suffixes."""
    product_path = Path(data_dir) / product_name / subcategory
    
    if not product_path.exists():
        raise FileNotFoundError(f"Product directory not found: {product_path}")
    
    # Define required suffixes
    required_suffixes = {
        'bullet_clusters': 'bullet_clusters.parquet',
        'bullet_diagram': 'bullet_diagram.parquet',
        'bullet_labels': 'bullet_labels.parquet',
        'bullet_keywords': 'bullet_keywords.parquet',
        'keyword_phrases': 'keyword_phrases.parquet'
    }
    
    data = {}
    
    # Find files matching each suffix
    for key, suffix in required_suffixes.items():
        matching_files = list(product_path.glob(f'*{suffix}'))
        
        if not matching_files:
            raise FileNotFoundError(f"No file ending with '{suffix}' found in {product_path}")
        elif len(matching_files) > 1:
            raise ValueError(f"Multiple files ending with '{suffix}' found in {product_path}: {matching_files}")
        
        data[key] = pd.read_parquet(matching_files[0])

    # Load secondary keywords from txt file
    secondary_keywords_files = list(product_path.glob('*secondary_keywords.txt'))
    
    if not secondary_keywords_files:
        raise FileNotFoundError(f"No file ending with 'secondary_keywords.txt' found in {product_path}")
    elif len(secondary_keywords_files) > 1:
        raise ValueError(f"Multiple files ending with 'secondary_keywords.txt' found in {product_path}: {secondary_keywords_files}")
    
    with open(secondary_keywords_files[0], 'r') as file:
        secondary_keywords = [line.strip() for line in file if line.strip()]
    
    return data, secondary_keywords

def count_label_occurrences(label_dict, bullet_labels, label):
    """
    Count how many times each subheading appears in bullet_labels 
    for a specific label category.
    
    Args:
        label_dict: Dictionary mapping labels to lists of subheadings
        bullet_labels: List of lists containing subheadings
        label: The label category to count (e.g., 'Durability')
    
    Returns:
        Dictionary mapping each subheading to its count
    
    Example:
        >>> count_label_occurrences(label_dict, bullet_labels, 'Durability')
        {'classic & durable': 15, 'built to last': 8, '100 handmade': 2, ...}
    """
    from collections import Counter
    bullet_labels_list_of_lists = bullet_labels.parsed_headings.values.tolist()
    # Get the list of subheadings for this label
    if label not in label_dict:
        return {}
    
    target_subheadings = set(label_dict[label])
    
    # Flatten bullet_labels and count occurrences
    all_subheadings = [item for sublist in bullet_labels_list_of_lists for item in sublist]
    
    # Count only the subheadings that belong to this label
    counts = dict(Counter(sub for sub in all_subheadings if sub in target_subheadings))
    counts_df = pd.DataFrame(list(counts.items()), columns=['subheading', 'count'])
    counts_df = counts_df.sort_values('count', ascending=False).reset_index(drop=True)
    
    return counts_df



# KEYWORD SUMMARIZER (after Gemini outputs listing)
def _plural_pattern(word: str) -> str:
    """
    Create a regex that matches a word and its common plural forms.
    """

    # words ending with 'y' → babies
    if word.endswith("y") and len(word) > 1 and word[-2] not in "aeiou":
        root = word[:-1]
        return rf"{re.escape(root)}(?:y|ies)"

    # words ending with s, x, z, ch, sh → boxes, dishes
    elif word.endswith(("s", "x", "z", "ch", "sh")):
        return rf"{re.escape(word)}(?:es)?"

    # default → frames, weddings
    else:
        return rf"{re.escape(word)}s?"
    

def keyword_count_df(keywords: list[str], text: str) -> pd.DataFrame:
    """
    Counts occurrences of each keyword in text and returns a dataframe.

    - Case insensitive
    - Matches whole words/phrases
    - Supports plural forms (wedding == weddings)
    """

    text_lower = text.lower()
    results = {}

    for kw in keywords:
        kw_clean = kw.strip().lower()
        if not kw_clean:
            continue

        # build plural-aware phrase pattern
        words = kw_clean.split()
        word_patterns = [_plural_pattern(w) for w in words]
        phrase_pattern = r"\b" + r"\s+".join(word_patterns) + r"\b"

        matches = re.findall(phrase_pattern, text_lower)
        results[kw_clean] = len(matches)

    df = pd.DataFrame(
        list(results.items()),
        columns=["keyword", "count"]
    ).sort_values(by="count", ascending=False).reset_index(drop=True)

    return df[df["count"] > 0]

def summarize_listing_keyword_stats(keywords: list[str], text: str):
    df = keyword_count_df(keywords, text)
    total_keywords_unique = len(df.keyword.unique())

    return total_keywords_unique, df




def filter_by_phrase(input_phrase, df, length=None):
    """
    Filter dataframe rows where 'search_term' contains all words from input_phrase.
    
    Args:
        input_phrase: String containing words to search for
        df: DataFrame with 'search_term' column
        length: Minimum number of words required in search_term (optional)
        
    Returns:
        Filtered DataFrame
    """
    # Normalize the input phrase: lowercase and handle dimension patterns
    input_phrase_lower = input_phrase.lower()
    
    # Extract words from input phrase
    words = input_phrase_lower.split()
    
    # Create regex patterns for each word
    patterns = []
    for word in words:
        # Check if word is a dimension pattern (e.g., "5x7")
        if re.match(r'\d+x\d+', word):
            # Make spaces around 'x' optional: 5x7, 5 x7, 5 x 7 all match
            dimension_pattern = word.replace('x', r'\s*x\s*')
            patterns.append(dimension_pattern)
        else:
            # Regular word - just escape special regex characters
            patterns.append(re.escape(word))
    
    # Filter function to apply to each row
    def contains_all_words(search_term):
        if pd.isna(search_term):
            return False
        
        search_term_lower = str(search_term).lower()
        
        # Check if all patterns are found in the search term
        for pattern in patterns:
            if not re.search(pattern, search_term_lower):
                return False
        return True
    
    # Count words function
    def count_words(search_term):
        if pd.isna(search_term):
            return 0
        
        search_term_str = str(search_term)
        # Normalize dimension patterns (5x7, 5 x7, 5 x 7) to single word
        normalized = re.sub(r'\d+\s*x\s*\d+', 'DIMENSION', search_term_str)
        # Split by whitespace and count
        words = normalized.split()
        return len(words)
    
    # Apply phrase filter
    filtered_df = df[df['search_term'].apply(contains_all_words)]
    
    # Apply length filter if specified
    if length is not None:
        filtered_df = filtered_df[filtered_df['search_term'].apply(count_words) >= length]
    
    return filtered_df



def filter_by_phrase_complex(df, input_phrase_and=None, input_phrase_or=None, length=None):
    """
    Filter dataframe rows based on keyword matching conditions.
    
    Args:
        df: DataFrame with 'search_term' column
        input_phrase_and: String containing words that ALL must be present (optional)
        input_phrase_or: String containing words where AT LEAST ONE must be present (optional)
        length: Minimum number of words required in search_term (optional)
        
    Returns:
        Filtered DataFrame
    """
    
    def create_patterns(phrase):
        """Helper function to create regex patterns from a phrase."""
        if not phrase:
            return []
        
        phrase_lower = phrase.lower()
        words = phrase_lower.split()
        patterns = []
        
        for word in words:
            # Check if word is a dimension pattern (e.g., "5x7")
            if re.match(r'\d+x\d+', word):
                # Make spaces around 'x' optional: 5x7, 5 x7, 5 x 7 all match
                dimension_pattern = word.replace('x', r'\s*x\s*')
                patterns.append(dimension_pattern)
            else:
                # Regular word - just escape special regex characters
                patterns.append(re.escape(word))
        
        return patterns
    
    # Create patterns for both phrase types
    patterns_and = create_patterns(input_phrase_and)
    patterns_or = create_patterns(input_phrase_or)
    
    # Filter function to apply to each row
    def matches_conditions(search_term):
        if pd.isna(search_term):
            return False
        
        search_term_lower = str(search_term).lower()
        
        # Check AND condition: all patterns must be found
        if patterns_and:
            for pattern in patterns_and:
                if not re.search(pattern, search_term_lower):
                    return False
        
        # Check OR condition: at least one pattern must be found
        if patterns_or:
            found_any = False
            for pattern in patterns_or:
                if re.search(pattern, search_term_lower):
                    found_any = True
                    break
            if not found_any:
                return False
        
        return True
    
    # Count words function
    def count_words(search_term):
        if pd.isna(search_term):
            return 0
        
        search_term_str = str(search_term)
        # Normalize dimension patterns (5x7, 5 x7, 5 x 7) to single word
        normalized = re.sub(r'\d+\s*x\s*\d+', 'DIMENSION', search_term_str)
        # Split by whitespace and count
        words = normalized.split()
        return len(words)
    
    # Apply phrase filters
    filtered_df = df[df['search_term'].apply(matches_conditions)]
    
    # Apply length filter if specified
    if length is not None:
        filtered_df = filtered_df[filtered_df['search_term'].apply(count_words) >= length]
    
    return filtered_df


capitalize_first = lambda lst: [re.sub(r'^([a-zA-Z])', lambda m: m.group(1).upper(), s) for s in lst]


def normalize_dimensions(text):
    """
    Normalize dimension patterns like '5x7', '5 x 7', '8x10', etc. to a standard format.
    Also normalizes measurements, weights, volumes, and other product specifications.
    """
    # Normalize spaced dimensions: '5 x 7' -> '5x7'
    text = re.sub(r'(\d+)\s*x\s*(\d+)', r'\1x\2', text, flags=re.IGNORECASE)
    
    # Normalize fluid ounces: '20 oz', '20oz', '20 fl oz' -> '20oz'
    text = re.sub(r'(\d+)\s*fl\s*oz', r'\1oz', text, flags=re.IGNORECASE)
    text = re.sub(r'(\d+)\s*oz\b', r'\1oz', text, flags=re.IGNORECASE)
    
    # Normalize pounds: '5 lb', '5 lbs', '5lb' -> '5lb'
    text = re.sub(r'(\d+)\s*lbs?\b', r'\1lb', text, flags=re.IGNORECASE)
    
    # Normalize inches: '12 in', '12 inch', '12"' -> '12in'
    text = re.sub(r'(\d+)\s*(?:inch|in)\b', r'\1in', text, flags=re.IGNORECASE)
    text = re.sub(r'(\d+)\s*"', r'\1in', text)
    
    # Normalize feet: '6 ft', '6 feet', "6'" -> '6ft'
    text = re.sub(r'(\d+)\s*(?:feet|ft)\b', r'\1ft', text, flags=re.IGNORECASE)
    text = re.sub(r"(\d+)\s*'", r'\1ft', text)
    
    # Normalize milliliters: '500 ml', '500ml' -> '500ml'
    text = re.sub(r'(\d+)\s*ml\b', r'\1ml', text, flags=re.IGNORECASE)
    
    # Normalize liters: '2 l', '2 liter' -> '2l'
    text = re.sub(r'(\d+)\s*(?:liter|l)\b', r'\1l', text, flags=re.IGNORECASE)
    
    # Normalize grams: '100 g', '100g' -> '100g'
    text = re.sub(r'(\d+)\s*g\b', r'\1g', text, flags=re.IGNORECASE)
    
    # Normalize kilograms: '2 kg', '2kg' -> '2kg'
    text = re.sub(r'(\d+)\s*kg\b', r'\1kg', text, flags=re.IGNORECASE)
    
    # Normalize watts/volts: '100 w', '120 v' -> '100w', '120v'
    text = re.sub(r'(\d+)\s*w\b', r'\1w', text, flags=re.IGNORECASE)
    text = re.sub(r'(\d+)\s*v\b', r'\1v', text, flags=re.IGNORECASE)
    
    # Normalize pack/count: '12 pack', '12-pack' -> '12pack'
    text = re.sub(r'(\d+)\s*-?\s*pack\b', r'\1pack', text, flags=re.IGNORECASE)
    text = re.sub(r'(\d+)\s*-?\s*count\b', r'\1count', text, flags=re.IGNORECASE)
    
    # Normalize percentages: '50 %' -> '50%'
    text = re.sub(r'(\d+)\s*%', r'\1%', text)
    
    return text

def filter_unique_keyword_phrases(df, search_term_col='search_term', keyword_rank_by='monthly_searches', ignore_plural=True):
    """
    Filter search terms to keep only phrases with unique keywords.
    
    For phrases with the same words in different orders, keeps the one with highest monthly_searches.
    Treats spaced dimensions (5 x 7) same as compact (5x7).
    
    Args:
        df: DataFrame with search terms and monthly_searches
        search_term_col: Column name for search terms
        keyword_rank_by: Column name for monthly_searches
        ignore_plural: If True, treat 'frame' and 'frames' as same word
    
    Returns:
        Filtered DataFrame with only phrases containing new unique words
    """
    # Sort by monthly_searches descending to prioritize high-impression phrases
    df_sorted = df.sort_values(keyword_rank_by, ascending=False).reset_index(drop=True)
    
    seen_words = set()
    keep_indices = []
    
    for idx, row in df_sorted.iterrows():
        phrase = row[search_term_col].lower().strip()
        
        # Normalize dimensions first
        phrase = normalize_dimensions(phrase)
        
        words = phrase.split()
        
        # Normalize words (handle plurals)
        normalized_words = set()
        for word in words:
            if ignore_plural and word.endswith('s') and len(word) > 1:
                # Try singular form
                normalized_words.add(word[:-1])
            else:
                normalized_words.add(word)
        
        # Check if this phrase introduces any new words
        new_words = normalized_words - seen_words
        
        if new_words:
            # This phrase has at least one new word - keep it
            keep_indices.append(idx)
            seen_words.update(normalized_words)
    
    # Return filtered dataframe in original impression order
    result_df = df_sorted.iloc[keep_indices].reset_index(drop=True)
    
    return result_df


def filter_by_phrase_final(df, input_phrase_and=None, input_phrase_or=None,
                           search_term_col='search_term', keyword_rank_by='monthly_searches', ignore_plural=True):

    all_keywords = filter_by_phrase_complex(df, input_phrase_and, input_phrase_or).sort_values(by = keyword_rank_by, ascending = False)
    res = filter_unique_keyword_phrases(all_keywords, search_term_col, keyword_rank_by, ignore_plural)
    return res
    

def sort_search_terms(search_terms, df, rank_by='monthly_searches'):
    # Validate that the rank_by column exists in df
    if rank_by not in df.columns:
        raise ValueError(f"Column '{rank_by}' not found in dataframe. Available columns: {df.columns.tolist()}")
    
    # Create a dictionary mapping search_term to the ranking metric
    ranking_dict = df.set_index('search_term')[rank_by].to_dict()
    
    # Sort the search_terms list
    sorted_terms = sorted(
        search_terms,
        key=lambda x: (x not in ranking_dict, -ranking_dict.get(x, 0))
    )
    
    return sorted_terms



def normalize_word(word, plural = True):
    """
    Normalize a word for comparison:
    - Remove plural 's' (frames -> frame)
    """
    # Handle plurals: remove trailing 's' if word ends with 's'
    if plural and len(word) > 3 and word.endswith('s') and not word.endswith('ss'):
        return word[:-1]
    
    return word

def extract_unique_words(search_terms, stop_words = STOP_WORDS, plural = True):
    """
    Extract unique words from search terms in order of first appearance.
    Handles plurals and format variations (e.g., 5x7, 5 x 7, 5  x 7).
    
    Parameters:
    -----------
    search_terms : list
        List of search term strings
    
    Returns:
    --------
    list
        List of unique words in order of first appearance
    """

    unique_words = []
    seen_normalized = set()
    for term in search_terms:
        # Split into words
        words = term.lower().split()
        
        for word in words:
            if word in stop_words:
                continue
            normalized = normalize_word(word, plural = plural)
            
            # If we haven't seen this normalized form, add the original word
            if normalized not in seen_normalized:
                seen_normalized.add(normalized)
                unique_words.append(word)
    
    return unique_words


def get_unique_words_from_string(text, max_chars=250, plural = False):
    cleaned = re.sub(r'[^\w\s]', '', text)
    unique_words = extract_unique_words([cleaned], plural = plural)
    
    result = []
    total_chars = 0
    for word in unique_words:
        # +1 for the space, except for the first word
        word_len = len(word) + (1 if result else 0)
        if total_chars + word_len > max_chars:
            break
        result.append(word)
        total_chars += word_len
    
    return " ".join(result)


def get_normalized_words(phrase, stop_words=STOP_WORDS):
    """
    Extract normalized words from a phrase (excluding stop words).
    
    Parameters:
    -----------
    phrase : str
        Search phrase to process
    stop_words : set
        Set of stop words to exclude
    
    Returns:
    --------
    set
        Set of normalized words from the phrase
    """
    # Normalize dimensions in the entire phrase
    normalized_phrase = normalize_dimensions(phrase.lower())
    
    # Split into words
    words = normalized_phrase.split()
    
    normalized_words = set()
    for word in words:
        # Skip stop words
        if word in stop_words:
            continue
        
        # Normalize for plurals
        normalized = normalize_word(word)
        normalized_words.add(normalized)
    
    return normalized_words

def get_top_n_search_terms(search_phrases, n=3, stop_words=STOP_WORDS):
    """
    Get the first n search phrases that each contribute at least 1 new word.
    Handles plurals (frame/frames) and dimension formats (5x7/5 x 7) as identical.
    
    Parameters:
    -----------
    search_phrases : list
        List of search phrase strings
    n : int
        Number of phrases to return (default: 3)
    stop_words : set
        Set of stop words to exclude (default: STOP_WORDS)
    
    Returns:
    --------
    list
        List of up to n search phrases that contribute new words
    """
    selected_phrases = []
    seen_words = set()
    
    for phrase in search_phrases:
        # Get normalized words from this phrase
        phrase_words = get_normalized_words(phrase, stop_words)
        
        # Check if this phrase contributes at least 1 new word
        new_words = phrase_words - seen_words
        
        if new_words:
            # This phrase contributes at least 1 new word
            selected_phrases.append(phrase)
            seen_words.update(phrase_words)
            
            # Stop if we've collected n phrases
            if len(selected_phrases) == n:
                break
    
    return selected_phrases



# Bullet Grid to Outline
def create_bullet_outline(bullet_grid, diff_threshold = 0.08):
    """
    Create Amazon listing bullet outline based on topic distribution.
    
    Parameters:
    -----------
    bullet_grid : pd.DataFrame
        DataFrame with topics as index and columns bullet_1 through bullet_5
        
    Returns:
    --------
    dict : Dictionary mapping bullet number to list of assigned topics
    """
    import pandas as pd
    
    outline = {}
    assigned_topics = set()
    columns = bullet_grid.columns.tolist()
    
    # Rule 1 (Enhanced): If top topic is > diff_threshold above second, it's a candidate
    # If a topic is top for multiple columns, assign to column with highest value
    rule1_candidates = {}  # topic -> [(col, value), ...]
    
    for col in columns:
        sorted_vals = bullet_grid[col].sort_values(ascending=False)
        if len(sorted_vals) >= 2:
            diff = sorted_vals.iloc[0] - sorted_vals.iloc[1]
            if diff > diff_threshold:
                topic = sorted_vals.index[0]
                value = sorted_vals.iloc[0]
                if topic not in rule1_candidates:
                    rule1_candidates[topic] = []
                rule1_candidates[topic].append((col, value))
    
    # Assign each topic to its highest-value column
    assigned_columns = set()
    for topic, col_vals in rule1_candidates.items():
        col_vals.sort(key=lambda x: x[1], reverse=True)
        max_col = col_vals[0][0]
        bullet_num = int(max_col.split('_')[1])
        outline[bullet_num] = [topic]
        assigned_topics.add(topic)
        assigned_columns.add(max_col)
        
        # For other columns where this was top, assign their second topic
        # (only if it also meets Rule 1 criteria)
        for col, _ in col_vals[1:]:
            sorted_vals = bullet_grid[col].sort_values(ascending=False)
            # Second topic is now the effective "top" for this column
            if len(sorted_vals) >= 2:
                second_topic = sorted_vals.index[1]
                # Check if second topic is > 0.08 above third
                if len(sorted_vals) >= 3:
                    second_val = sorted_vals.iloc[1]
                    third_val = sorted_vals.iloc[2]
                    if (second_val - third_val) > 0.08 and second_topic not in assigned_topics:
                        bullet_num = int(col.split('_')[1])
                        outline[bullet_num] = [second_topic]
                        assigned_topics.add(second_topic)
                        assigned_columns.add(col)
    
    # Get remaining columns
    remaining_cols = [col for col in columns if col not in assigned_columns]
    
    # Rules 2-3: Iteratively resolve conflicts
    max_iterations = 10
    for iteration in range(max_iterations):
        if not remaining_cols:
            break
            
        # Get top 2 unassigned topics for each remaining column
        top_2_by_col = {}
        for col in remaining_cols:
            sorted_vals = bullet_grid[col].sort_values(ascending=False)
            top_2 = [topic for topic in sorted_vals.index 
                    if topic not in assigned_topics][:2]
            top_2_by_col[col] = top_2
        
        # Find which topics appear in multiple columns' top 2
        topic_to_cols = {}
        for col, topics in top_2_by_col.items():
            for topic in topics:
                if topic not in topic_to_cols:
                    topic_to_cols[topic] = []
                topic_to_cols[topic].append((col, bullet_grid.loc[topic, col]))
        
        # Filter to only multi-column topics and sort by max value
        multi_col_topics = [(topic, col_vals) for topic, col_vals in topic_to_cols.items() 
                           if len(col_vals) > 1]
        
        if not multi_col_topics:
            break
        
        # Sort by maximum value across all columns (descending)
        multi_col_topics.sort(key=lambda x: max(val for _, val in x[1]), reverse=True)
        
        # Process the topic with highest max value
        topic, col_vals = multi_col_topics[0]
        
        # Sort by value and assign to highest
        col_vals.sort(key=lambda x: x[1], reverse=True)
        max_col = col_vals[0][0]
        bullet_num = int(max_col.split('_')[1])
        outline[bullet_num] = [topic]
        assigned_topics.add(topic)
        remaining_cols.remove(max_col)
        
        # For other columns where this topic was in top 2,
        # assign the OTHER topic from their top 2
        for col, val in col_vals[1:]:
            if col in remaining_cols:
                other_topics = [t for t in top_2_by_col[col] 
                               if t != topic and t not in assigned_topics]
                if other_topics:
                    other_topic = other_topics[0]
                    bullet_num = int(col.split('_')[1])
                    outline[bullet_num] = [other_topic]
                    assigned_topics.add(other_topic)
                    remaining_cols.remove(col)
                    
                    # CASCADING: Check if this newly assigned topic appears 
                    # in OTHER remaining columns' top 2
                    for other_col in remaining_cols[:]:
                        if other_topic in top_2_by_col.get(other_col, []):
                            # Assign the OTHER topic from that column's top 2
                            cascade_topics = [t for t in top_2_by_col[other_col]
                                            if t != other_topic and t not in assigned_topics]
                            if cascade_topics:
                                cascade_topic = cascade_topics[0]
                                bullet_num = int(other_col.split('_')[1])
                                outline[bullet_num] = [cascade_topic]
                                assigned_topics.add(cascade_topic)
                                remaining_cols.remove(other_col)
    
    # Rule 4: Remaining columns keep their top 2 unassigned topics
    for col in remaining_cols:
        sorted_vals = bullet_grid[col].sort_values(ascending=False)
        top_2 = [topic for topic in sorted_vals.index 
                if topic not in assigned_topics][:2]
        bullet_num = int(col.split('_')[1])
        outline[bullet_num] = top_2
    
    # Sort by bullet number and return
    return {k: outline[k] for k in sorted(outline.keys())}


def extract_amazon_asin(url: str) -> Optional[str]:
    """
    Extract Amazon ASIN from a URL.
    
    Args:
        url: Amazon product URL (full or partial)
        
    Returns:
        10-character ASIN string if found, None otherwise

    """
    if not url or not isinstance(url, str):
        return None
    
    # Clean the URL
    url = url.strip()
    
    # Pattern 1: /dp/{ASIN} - most common
    match = re.search(r'/dp/([A-Z0-9]{10})(?:[/?]|$)', url, re.IGNORECASE)
    if match:
        return match.group(1).upper()
    
    # Pattern 2: /gp/product/{ASIN} - alternate format
    match = re.search(r'/gp/product/([A-Z0-9]{10})(?:[/?]|$)', url, re.IGNORECASE)
    if match:
        return match.group(1).upper()
    
    # Pattern 3: /d/{ASIN} - short format
    match = re.search(r'/d/([A-Z0-9]{10})(?:[/?]|$)', url, re.IGNORECASE)
    if match:
        return match.group(1).upper()
    
    # Pattern 4: Query parameter ?asin={ASIN}
    try:
        parsed = urlparse(url)
        query_params = parse_qs(parsed.query)
        if 'asin' in query_params:
            asin = query_params['asin'][0]
            if len(asin) == 10 and re.match(r'^[A-Z0-9]{10}$', asin, re.IGNORECASE):
                return asin.upper()
    except Exception:
        pass
    
    # Pattern 5: Fallback - any 10-char alphanumeric after /product/ or in path
    match = re.search(r'/product/([A-Z0-9]{10})(?:[/?]|$)', url, re.IGNORECASE)
    if match:
        return match.group(1).upper()
    
    # Pattern 6: Last resort - find any standalone 10-char alphanumeric 
    # that looks like an ASIN (starts with B0 or is all digits)
    matches = re.findall(r'\b([A-Z0-9]{10})\b', url, re.IGNORECASE)
    for potential_asin in matches:
        # Prioritize ASINs starting with B0 (most products)
        if potential_asin.upper().startswith('B0'):
            return potential_asin.upper()
    
    # If found any 10-char alphanumeric, return the first one
    if matches:
        return matches[0].upper()
    
    return None











def create_listing_docx(title, listing_content):
    """
    Create a Word document matching the template format.
    Returns the document as bytes.
    """
    doc = Document()
    
    # Add the title (bold)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = 'Calibri'
    
    # Add empty line
    doc.add_paragraph()
    
    # Parse and add bullet points
    lines = listing_content.strip().split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if this is a bullet point
        if line.startswith('•') or line.startswith('-') or line.startswith('*'):
            # Remove the bullet character
            content = line.lstrip('•-* ').strip()
            
            # Split on the first colon to get bold header and description
            if ':' in content:
                parts = content.split(':', 1)
                header = parts[0].strip()
                description = parts[1].strip()
                
                # Create bullet paragraph
                bullet_para = doc.add_paragraph(style='List Bullet')
                
                # Add bold header
                header_run = bullet_para.add_run(header + ':')
                header_run.bold = True
                header_run.font.size = Pt(11)
                header_run.font.name = 'Calibri'
                
                # Add space
                bullet_para.add_run(' ')
                
                # Add description
                desc_run = bullet_para.add_run(description)
                desc_run.font.size = Pt(11)
                desc_run.font.name = 'Calibri'
            else:
                # Just add as bullet without special formatting
                bullet_para = doc.add_paragraph(content, style='List Bullet')
                for run in bullet_para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
    
    # Add empty line
    doc.add_paragraph()
    
    # Add date
    date_para = doc.add_paragraph()
    current_date = datetime.now().strftime('%-m/%-d/%Y %-I:%M%p').lower()
    date_run = date_para.add_run(f'Date: {current_date}')
    date_run.font.size = Pt(11)
    date_run.font.name = 'Calibri'
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()